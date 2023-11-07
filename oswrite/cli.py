import os
import sys
import click
from openai import OpenAI
from docx import Document


def transcribe(audio_file, api_key=None):
    client = OpenAI(api_key=api_key)
    audio_file = open(audio_file, "rb")
    transcript = client.audio.transcriptions.create(
        model="whisper-1", file=audio_file, response_format="text"
    )
    return transcript


@click.command()
@click.version_option()
@click.option("--audio_file", help="Audio file path", required=True)
@click.option("--output", help="Output path to save transcribed text")
@click.option(
    "--token",
    help="OpenAI API key",
    envvar="OPENAI_API_KEY",
)
def cli(audio_file, output, token):
    base_file_name = os.path.basename(audio_file)
    if not (base_file_name.endswith(".mp3") or base_file_name.endswith(".wav")):
        raise click.BadOptionUsage(
            "audio_file", "Input file must be .mp3 or .wav format"
        )
    result = transcribe(audio_file, token)
    if output:
        file_extension = os.path.basename(output).split(".")[-1]
        if file_extension == "txt":
            f = open(output, "w")
            f.write(result)
        elif file_extension == "docx":
            document = Document()
            document.add_heading("Whisper Transcribed Text", level=1)
            document.add_paragraph(result)
            document.save(output)

    sys.stdout.write(str(result) + "\n")
    sys.stdout.flush()  # Ensure everything in the buffer is written out
